"""
Report Export System

Provides multi-format export capabilities for code review reports.
Supports PDF, Excel, Word, and other formats with charts and detailed analysis.
"""

import io
import json
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, BinaryIO
from dataclasses import dataclass, field

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.linecharts import HorizontalLineChart
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class ExportFormat(str, Enum):
    """Export format types"""
    PDF = "pdf"
    EXCEL = "excel"
    WORD = "word"
    HTML = "html"
    JSON = "json"
    CSV = "csv"


class ChartType(str, Enum):
    """Chart types for reports"""
    LINE = "line"
    BAR = "bar"
    PIE = "pie"
    TABLE = "table"


@dataclass
class ChartData:
    """Chart data structure"""
    type: ChartType
    title: str
    data: Dict[str, Any]
    labels: Optional[List[str]] = None
    colors: Optional[List[str]] = None


@dataclass
class ReportSection:
    """Report section structure"""
    title: str
    content: str
    charts: List[ChartData] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    subsections: List['ReportSection'] = field(default_factory=list)


@dataclass
class ExportReport:
    """Complete report structure"""
    title: str
    subtitle: Optional[str] = None
    author: Optional[str] = None
    date: Optional[datetime] = None
    summary: Optional[str] = None
    sections: List[ReportSection] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class PDFExporter:
    """PDF report exporter using ReportLab"""

    def __init__(self):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30
        ))
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12
        ))

    def export(self, report: ExportReport, output: BinaryIO) -> None:
        """Export report to PDF"""
        doc = SimpleDocTemplate(output, pagesize=letter)
        story = []

        # Title page
        story.append(Paragraph(report.title, self.styles['CustomTitle']))
        if report.subtitle:
            story.append(Paragraph(report.subtitle, self.styles['Heading2']))
        if report.author:
            story.append(Paragraph(f"Author: {report.author}", self.styles['Normal']))
        if report.date:
            story.append(Paragraph(f"Date: {report.date.strftime('%Y-%m-%d %H:%M')}", self.styles['Normal']))
        story.append(Spacer(1, 0.5*inch))

        # Summary
        if report.summary:
            story.append(Paragraph("Executive Summary", self.styles['CustomHeading']))
            story.append(Paragraph(report.summary, self.styles['Normal']))
            story.append(Spacer(1, 0.3*inch))

        # Sections
        for section in report.sections:
            self._add_section(story, section)

        doc.build(story)

    def _add_section(self, story: List, section: ReportSection, level: int = 1):
        """Add a section to the PDF"""
        # Section title
        style = self.styles['CustomHeading'] if level == 1 else self.styles['Heading3']
        story.append(Paragraph(section.title, style))

        # Section content
        if section.content:
            story.append(Paragraph(section.content, self.styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

        # Charts
        for chart_data in section.charts:
            chart = self._create_chart(chart_data)
            if chart:
                story.append(chart)
                story.append(Spacer(1, 0.2*inch))

        # Tables
        for table_data in section.tables:
            table = self._create_table(table_data)
            if table:
                story.append(table)
                story.append(Spacer(1, 0.2*inch))

        # Subsections
        for subsection in section.subsections:
            self._add_section(story, subsection, level + 1)

    def _create_chart(self, chart_data: ChartData) -> Optional[Drawing]:
        """Create a chart drawing"""
        drawing = Drawing(400, 200)

        if chart_data.type == ChartType.LINE:
            chart = HorizontalLineChart()
            chart.x = 50
            chart.y = 50
            chart.height = 125
            chart.width = 300
            chart.data = [list(chart_data.data.values())]
            chart.categoryAxis.categoryNames = list(chart_data.data.keys())
            drawing.add(chart)

        elif chart_data.type == ChartType.BAR:
            chart = VerticalBarChart()
            chart.x = 50
            chart.y = 50
            chart.height = 125
            chart.width = 300
            chart.data = [list(chart_data.data.values())]
            chart.categoryAxis.categoryNames = list(chart_data.data.keys())
            drawing.add(chart)

        elif chart_data.type == ChartType.PIE:
            chart = Pie()
            chart.x = 150
            chart.y = 65
            chart.width = 100
            chart.height = 100
            chart.data = list(chart_data.data.values())
            chart.labels = list(chart_data.data.keys())
            drawing.add(chart)

        return drawing

    def _create_table(self, table_data: Dict[str, Any]) -> Optional[Table]:
        """Create a table"""
        if 'headers' not in table_data or 'rows' not in table_data:
            return None

        data = [table_data['headers']] + table_data['rows']
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        return table


class ExcelExporter:
    """Excel report exporter using openpyxl"""

    def __init__(self):
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel export")

    def export(self, report: ExportReport, output: BinaryIO) -> None:
        """Export report to Excel"""
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet

        # Summary sheet
        ws_summary = wb.create_sheet("Summary")
        self._add_summary_sheet(ws_summary, report)

        # Section sheets
        for idx, section in enumerate(report.sections):
            ws = wb.create_sheet(section.title[:31])  # Excel sheet name limit
            self._add_section_sheet(ws, section)

        wb.save(output)

    def _add_summary_sheet(self, ws, report: ExportReport):
        """Add summary sheet"""
        # Title
        ws['A1'] = report.title
        ws['A1'].font = Font(size=18, bold=True)
        ws.merge_cells('A1:D1')

        row = 3
        if report.subtitle:
            ws[f'A{row}'] = report.subtitle
            ws[f'A{row}'].font = Font(size=14)
            row += 1

        if report.author:
            ws[f'A{row}'] = f"Author: {report.author}"
            row += 1

        if report.date:
            ws[f'A{row}'] = f"Date: {report.date.strftime('%Y-%m-%d %H:%M')}"
            row += 2

        if report.summary:
            ws[f'A{row}'] = "Executive Summary"
            ws[f'A{row}'].font = Font(size=12, bold=True)
            row += 1
            ws[f'A{row}'] = report.summary
            ws.merge_cells(f'A{row}:D{row}')

    def _add_section_sheet(self, ws, section: ReportSection):
        """Add section sheet"""
        row = 1

        # Section title
        ws[f'A{row}'] = section.title
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 2

        # Content
        if section.content:
            ws[f'A{row}'] = section.content
            ws.merge_cells(f'A{row}:D{row}')
            row += 2

        # Tables
        for table_data in section.tables:
            row = self._add_table(ws, table_data, row)
            row += 2

        # Charts
        for chart_data in section.charts:
            row = self._add_chart(ws, chart_data, row)
            row += 2

    def _add_table(self, ws, table_data: Dict[str, Any], start_row: int) -> int:
        """Add table to worksheet"""
        if 'headers' not in table_data or 'rows' not in table_data:
            return start_row

        # Headers
        for col_idx, header in enumerate(table_data['headers'], 1):
            cell = ws.cell(row=start_row, column=col_idx, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        # Rows
        for row_idx, row_data in enumerate(table_data['rows'], start_row + 1):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        return start_row + len(table_data['rows']) + 1

    def _add_chart(self, ws, chart_data: ChartData, start_row: int) -> int:
        """Add chart to worksheet"""
        # Write data
        ws.cell(row=start_row, column=1, value=chart_data.title)
        ws.cell(row=start_row, column=1).font = Font(bold=True)

        data_row = start_row + 1
        for idx, (label, value) in enumerate(chart_data.data.items(), 1):
            ws.cell(row=data_row + idx, column=1, value=label)
            ws.cell(row=data_row + idx, column=2, value=value)

        # Create chart
        if chart_data.type == ChartType.BAR:
            chart = BarChart()
        elif chart_data.type == ChartType.LINE:
            chart = LineChart()
        elif chart_data.type == ChartType.PIE:
            chart = PieChart()
        else:
            return data_row + len(chart_data.data) + 1

        chart.title = chart_data.title
        data = Reference(ws, min_col=2, min_row=data_row + 1, max_row=data_row + len(chart_data.data))
        cats = Reference(ws, min_col=1, min_row=data_row + 1, max_row=data_row + len(chart_data.data))
        chart.add_data(data)
        chart.set_categories(cats)

        ws.add_chart(chart, f'D{start_row}')

        return data_row + len(chart_data.data) + 15  # Chart height


class WordExporter:
    """Word document exporter using python-docx"""

    def __init__(self):
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")

    def export(self, report: ExportReport, output: BinaryIO) -> None:
        """Export report to Word"""
        doc = Document()

        # Title
        title = doc.add_heading(report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if report.subtitle:
            subtitle = doc.add_heading(report.subtitle, 2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if report.author or report.date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if report.author:
                p.add_run(f"Author: {report.author}\n")
            if report.date:
                p.add_run(f"Date: {report.date.strftime('%Y-%m-%d %H:%M')}")

        doc.add_page_break()

        # Summary
        if report.summary:
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(report.summary)

        # Sections
        for section in report.sections:
            self._add_section(doc, section)

        doc.save(output)

    def _add_section(self, doc: Document, section: ReportSection, level: int = 1):
        """Add section to document"""
        doc.add_heading(section.title, level)

        if section.content:
            doc.add_paragraph(section.content)

        # Tables
        for table_data in section.tables:
            self._add_table(doc, table_data)

        # Subsections
        for subsection in section.subsections:
            self._add_section(doc, subsection, level + 1)

    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """Add table to document"""
        if 'headers' not in table_data or 'rows' not in table_data:
            return

        table = doc.add_table(rows=1 + len(table_data['rows']), cols=len(table_data['headers']))
        table.style = 'Light Grid Accent 1'

        # Headers
        for idx, header in enumerate(table_data['headers']):
            cell = table.rows[0].cells[idx]
            cell.text = str(header)
            cell.paragraphs[0].runs[0].font.bold = True

        # Rows
        for row_idx, row_data in enumerate(table_data['rows'], 1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(value)


class ReportExporter:
    """Main report exporter with multi-format support"""

    def __init__(self):
        self.exporters = {
            ExportFormat.PDF: PDFExporter() if REPORTLAB_AVAILABLE else None,
            ExportFormat.EXCEL: ExcelExporter() if OPENPYXL_AVAILABLE else None,
            ExportFormat.WORD: WordExporter() if PYTHON_DOCX_AVAILABLE else None,
        }

    def export(self, report: ExportReport, format: ExportFormat, output_path: Optional[str] = None) -> bytes:
        """
        Export report to specified format

        Args:
            report: Report to export
            format: Export format
            output_path: Optional output file path

        Returns:
            Exported report as bytes
        """
        if format == ExportFormat.JSON:
            return self._export_json(report, output_path)
        elif format == ExportFormat.CSV:
            return self._export_csv(report, output_path)
        elif format == ExportFormat.HTML:
            return self._export_html(report, output_path)

        exporter = self.exporters.get(format)
        if not exporter:
            raise ValueError(f"Exporter for format {format} is not available")

        buffer = io.BytesIO()
        exporter.export(report, buffer)
        data = buffer.getvalue()

        if output_path:
            Path(output_path).write_bytes(data)

        return data

    def _export_json(self, report: ExportReport, output_path: Optional[str] = None) -> bytes:
        """Export to JSON"""
        data = {
            'title': report.title,
            'subtitle': report.subtitle,
            'author': report.author,
            'date': report.date.isoformat() if report.date else None,
            'summary': report.summary,
            'sections': [self._section_to_dict(s) for s in report.sections],
            'metadata': report.metadata
        }
        json_str = json.dumps(data, indent=2)
        json_bytes = json_str.encode('utf-8')

        if output_path:
            Path(output_path).write_bytes(json_bytes)

        return json_bytes

    def _export_csv(self, report: ExportReport, output_path: Optional[str] = None) -> bytes:
        """Export to CSV (tables only)"""
        lines = [f"# {report.title}\n"]

        for section in report.sections:
            lines.append(f"\n## {section.title}\n")
            for table_data in section.tables:
                if 'headers' in table_data and 'rows' in table_data:
                    lines.append(','.join(str(h) for h in table_data['headers']) + '\n')
                    for row in table_data['rows']:
                        lines.append(','.join(str(v) for v in row) + '\n')
                    lines.append('\n')

        csv_str = ''.join(lines)
        csv_bytes = csv_str.encode('utf-8')

        if output_path:
            Path(output_path).write_bytes(csv_bytes)

        return csv_bytes

    def _export_html(self, report: ExportReport, output_path: Optional[str] = None) -> bytes:
        """Export to HTML"""
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="UTF-8">',
            f'<title>{report.title}</title>',
            '<style>',
            'body { font-family: Arial, sans-serif; margin: 40px; }',
            'h1 { color: #333; }',
            'h2 { color: #666; margin-top: 30px; }',
            'table { border-collapse: collapse; width: 100%; margin: 20px 0; }',
            'th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }',
            'th { background-color: #f2f2f2; }',
            '</style>',
            '</head>',
            '<body>',
            f'<h1>{report.title}</h1>',
        ]

        if report.subtitle:
            html_parts.append(f'<h2>{report.subtitle}</h2>')

        if report.summary:
            html_parts.append(f'<p><strong>Summary:</strong> {report.summary}</p>')

        for section in report.sections:
            html_parts.extend(self._section_to_html(section))

        html_parts.extend(['</body>', '</html>'])

        html_str = '\n'.join(html_parts)
        html_bytes = html_str.encode('utf-8')

        if output_path:
            Path(output_path).write_bytes(html_bytes)

        return html_bytes

    def _section_to_dict(self, section: ReportSection) -> Dict:
        """Convert section to dictionary"""
        return {
            'title': section.title,
            'content': section.content,
            'charts': [{'type': c.type, 'title': c.title, 'data': c.data} for c in section.charts],
            'tables': section.tables,
            'subsections': [self._section_to_dict(s) for s in section.subsections]
        }

    def _section_to_html(self, section: ReportSection, level: int = 2) -> List[str]:
        """Convert section to HTML"""
        parts = [f'<h{level}>{section.title}</h{level}>']

        if section.content:
            parts.append(f'<p>{section.content}</p>')

        for table_data in section.tables:
            if 'headers' in table_data and 'rows' in table_data:
                parts.append('<table>')
                parts.append('<tr>')
                for header in table_data['headers']:
                    parts.append(f'<th>{header}</th>')
                parts.append('</tr>')
                for row in table_data['rows']:
                    parts.append('<tr>')
                    for value in row:
                        parts.append(f'<td>{value}</td>')
                    parts.append('</tr>')
                parts.append('</table>')

        for subsection in section.subsections:
            parts.extend(self._section_to_html(subsection, level + 1))

        return parts

    def is_format_available(self, format: ExportFormat) -> bool:
        """Check if export format is available"""
        if format in [ExportFormat.JSON, ExportFormat.CSV, ExportFormat.HTML]:
            return True
        return self.exporters.get(format) is not None

    def get_available_formats(self) -> List[ExportFormat]:
        """Get list of available export formats"""
        formats = [ExportFormat.JSON, ExportFormat.CSV, ExportFormat.HTML]
        for fmt, exporter in self.exporters.items():
            if exporter:
                formats.append(fmt)
        return formats
